#!/usr/bin/env python3
"""Generate Cast It architecture documentation as a Word file with embedded diagrams."""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
DIAGRAMS = ROOT / "docs" / "architecture" / "diagrams"
OUTPUT = ROOT / "docs" / "Cast-It-Architecture.docx"

SECTIONS: list[tuple[str, str, str]] = [
    (
        "1. System Context",
        "01-system-context",
        (
            "High-level view of Cast It and its external dependencies. "
            "Staff use the Operations UI; API clients integrate via REST. "
            "The platform runs in Docker (web, Celery worker/beat, PostgreSQL, Redis) "
            "and connects to Ollama, Chatterbox TTS, news feeds, FFmpeg, and YouTube."
        ),
    ),
    (
        "2. Layered Architecture",
        "02-layered-architecture",
        (
            "Clean architecture separation: presentation (Operations, Admin, API) "
            "calls application services in services/, which use domain types and "
            "infrastructure adapters for external I/O. Django models in apps/ persist data."
        ),
    ),
    (
        "3. Content Pipeline",
        "03-content-pipeline",
        (
            "End-to-end podcast production flow: NewsSource → Article → Episode → "
            "Script/ScriptSegment → AudioAsset → PipelineRun → PublishedEpisode/RSS feed."
        ),
    ),
    (
        "4. Async Job Backbone",
        "04-async-jobs",
        (
            "All long-running work is dispatched via JobDispatchService, tracked as Job "
            "records, and executed by Celery workers through registered handlers on "
            "dedicated queues (ingestion, llm, tts, audio, publishing, monitoring)."
        ),
    ),
    (
        "5. Django Apps & Data Ownership",
        "05-django-apps",
        (
            "Django apps own domain models. operations/ has no models — it is a staff "
            "dashboard only. scheduler/ tracks background jobs; knowledge/ supports RAG."
        ),
    ),
    (
        "6. Operations UI vs Admin vs API",
        "06-entry-points",
        (
            "Three entry points share JobDispatchService. Operations UI additionally uses "
            "facades under services/admin/ and polls GET /api/jobs/{id}/ for progress."
        ),
    ),
    (
        "7. Docker Deployment (Local Dev)",
        "07-docker-deployment",
        (
            "docker-compose.yml runs web, celery-worker, celery-beat, db (pgvector), "
            "and redis. Ollama and Chatterbox run on the host and are reached via "
            "host.docker.internal from containers."
        ),
    ),
    (
        "8. RAG Side Path",
        "08-rag-path",
        (
            "Optional knowledge-base path: articles are chunked and embedded via Ollama, "
            "stored in pgvector, and retrieved to enrich script generation context."
        ),
    ),
]

QUEUE_TABLE = [
    ("ingestion", "import_news", "NewsImportService"),
    ("llm", "summarize_article, classify_article, episode_planning, generate_script", "LLMService, ScriptGenerationService"),
    ("tts", "generate_audio", "AudioGenerationService"),
    ("audio", "run_audio_pipeline", "AudioPipelineService"),
    ("publishing", "publish_episode", "PublishService"),
    ("monitoring", "health_check, retry_job", "CeleryHealthService, JobService"),
]

SUMMARY_TABLE = [
    ("Business logic", "services/"),
    ("External I/O", "infrastructure/"),
    ("Pure types / contracts", "domain/"),
    ("Persistence", "apps/*/models.py"),
    ("Staff UI", "apps/operations/ + templates/operations/"),
    ("Machine API", "api/v1/"),
    ("Background work", "apps/scheduler/tasks/ + Celery"),
    ("Config", "config/settings/"),
]


def render_diagrams() -> None:
    """Render .mmd sources to PNG if images are missing."""
    missing = [p for p in DIAGRAMS.glob("*.mmd") if not p.with_suffix(".png").exists()]
    if not missing:
        return
    for mmd in missing:
        png = mmd.with_suffix(".png")
        subprocess.run(
            [
                "npx",
                "-y",
                "@mermaid-js/mermaid-cli@11",
                "-i",
                str(mmd),
                "-o",
                str(png),
                "-b",
                "transparent",
                "-w",
                "1400",
            ],
            check=True,
            cwd=DIAGRAMS,
        )


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str) -> None:
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(8)


def add_image(doc: Document, image_path: Path, width_inches: float = 6.5) -> None:
    if not image_path.exists():
        doc.add_paragraph(f"[Missing diagram: {image_path.name}]")
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    doc.add_paragraph()


def add_table(doc: Document, headers: tuple[str, ...], rows: list[tuple[str, ...]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
    for row_idx, row in enumerate(rows, start=1):
        cells = table.rows[row_idx].cells
        for col_idx, value in enumerate(row):
            cells[col_idx].text = value
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()

    title = doc.add_heading("Cast It — Architecture Overview", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_body(
        doc,
        "Production-ready Django platform for automated podcast production: "
        "news ingestion, episode planning, script generation (LLM), "
        "audio synthesis (TTS), FFmpeg post-processing, and publishing.",
    )

    add_heading(doc, "Tech Stack", 2)
    add_body(
        doc,
        "Python 3.13+, Django 5+, DRF · PostgreSQL 16 + pgvector · Redis · "
        "Celery + Beat · Ollama (LLM/embeddings) · Chatterbox (TTS) · FFmpeg · "
        "Docker Compose · Gunicorn/Nginx (production).",
    )

    for title_text, stem, description in SECTIONS:
        add_heading(doc, title_text, 1)
        add_body(doc, description)
        add_image(doc, DIAGRAMS / f"{stem}.png")

    add_heading(doc, "Celery Queue Reference", 1)
    add_table(doc, ("Queue", "Job Types", "Primary Service"), QUEUE_TABLE)

    add_heading(doc, "Layer Summary", 1)
    add_table(doc, ("Concern", "Location"), SUMMARY_TABLE)

    add_heading(doc, "Diagram Sources", 2)
    add_body(
        doc,
        f"Mermaid source files: {DIAGRAMS.relative_to(ROOT)}/. "
        "Re-run scripts/generate_architecture_docx.py to regenerate this document.",
    )

    return doc


def main() -> int:
    render_diagrams()
    DIAGRAMS.mkdir(parents=True, exist_ok=True)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
